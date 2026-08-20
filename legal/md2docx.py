"""将听证材料 Markdown 转换为符合公文习惯的 docx（仿宋正文／黑体标题／表格保留）。"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BODY_FONT = "仿宋_GB2312"
BODY_FONT_FALLBACK = "仿宋"
HEAD_FONT = "黑体"


def set_run_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_inline(paragraph, text, font, size, bold=False):
    """处理 **加粗** 标记，其余按普通文本输出。"""
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not seg:
            continue
        set_run_font(paragraph.add_run(seg), font, size, bold or i % 2 == 1)


def clean(text):
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def convert(src: Path, dst: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT_FALLBACK
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_FALLBACK)

    lines = src.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if set(line.strip()) == {"-"} and len(line.strip()) >= 3:
            i += 1
            continue

        # 表格
        if line.lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [
                [c.strip() for c in r.strip().strip("|").split("|")]
                for r in block
                if not re.fullmatch(r"\|[\s:|-]+\|", r.strip())
            ]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row[: len(rows[0])]):
                        para = table.cell(r, c).paragraphs[0]
                        add_inline(para, clean(cell), BODY_FONT_FALLBACK, 10.5, r == 0)
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level, text = len(heading.group(1)), clean(heading.group(2))
            para = doc.add_paragraph()
            para.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            )
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            add_inline(para, text, HEAD_FONT, {1: 18, 2: 15, 3: 13.5}[level], True)
            i += 1
            continue

        # 引用块
        if line.lstrip().startswith(">"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                block.append(lines[i].lstrip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            para.paragraph_format.right_indent = Pt(12)
            add_inline(para, clean("".join(b for b in block if b)), BODY_FONT_FALLBACK, 12)
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            block = [bullet.group(1).strip()]
            i += 1
            while i < len(lines) and lines[i].startswith("  ") and lines[i].strip():
                if re.match(r"^\s*[-*]\s", lines[i]):
                    break
                block.append(lines[i].strip())
                i += 1
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, clean("".join(block)), BODY_FONT_FALLBACK, 12)
            continue

        numbered = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if numbered:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            add_inline(
                para,
                f"{numbered.group(1)}. {clean(numbered.group(2))}",
                BODY_FONT_FALLBACK,
                12,
            )
            i += 1
            continue

        # 正文段落：合并被硬换行拆开的续行
        block = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or re.match(r"^(#{1,3}\s|[-*]\s|\d+\.\s|>|\|)", nxt):
                break
            if set(nxt) == {"-"} and len(nxt) >= 3:
                break
            block.append(nxt)
            i += 1
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(24)
        para.paragraph_format.line_spacing = 1.5
        add_inline(para, clean("".join(block)), BODY_FONT_FALLBACK, 12)

    doc.save(dst)
    print(f"生成 {dst}")


if __name__ == "__main__":
    for path in sys.argv[1:]:
        src = Path(path)
        convert(src, src.with_suffix(".docx"))
